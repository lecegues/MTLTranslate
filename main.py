# For Development Plan 2
# Should take the code from splitfirst.py 
# Only change: Use OpenAI API to fix grammar

# for openai
import time
import os
import openai

# for extracting text
import split

# for adding to word document
import docx

# use os library to get the API key from the environment variable
api_key = os.environ["OPENAI_API_KEY"]

# set the openai api key
openai.api_key = api_key

def fixText(text: str):
    '''
    Fixes the grammar and readability of the text using OpenAI's API
    '''
    prompt = "Fix the grammar and readability of this: " + text
    response = openai.Completion.create(engine="text-davinci-003",prompt=prompt,max_tokens=2000)
    answer = response["choices"][0]["text"]

    return answer

def addDoc(filename: str, header: str, paragraphs: list):
    document = docx.Document()

    document.add_heading(header,0)

    for i in range(len(paragraphs)):
        document.add_paragraph(paragraphs[i])

    document.save(filename)

def main(website):

    # Retrieve website URL from text file (1-liners)
    word_limit =  200 # hardcoded
    
    novel_info = split.scrapeSite(website, 200)
    novel_paragraphs = novel_info['text']
    header= novel_info['header']

    fixed_novel_paragraphs = []


    for text in novel_paragraphs:
        to_add = fixText(text)
        fixed_novel_paragraphs.append(to_add)
        time.sleep(10)

    # now we just put the text into a word document
    # @TODO currently hard-coded - fix
    addDoc(header+".docx",header,fixed_novel_paragraphs)



    print("-----Ending Process-----")

if __name__ == "__main__":

    with open('sites.txt','r') as f:
        lines = f.readlines()
    
    for website in lines:
        main(website.strip())
        time.sleep(15)