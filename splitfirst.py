# John Lecegues
# Note: Currently coded for only NovelHall websites

# for openai-chat
import os
import openai

# for extracting text
from selenium import webdriver
from selenium.webdriver.common.by import By

# for extracting into word document
import docx
from docx import *

# initialize webdriver
driver = webdriver.Chrome()

# use os library to get the API key from the environment variable
api_key = os.environ["OPENAI_API_KEY"]

# set the openai api key
openai.api_key = api_key

def split_text_into_paragraphs(text: str, word_count: int) -> list:
    '''
    Should iterate through text until the given word count. From then, keep going until
    reaching the next period (.) or the end of the text (null)
    text (Str): full text to split into paragraphs.
    word_count (int): specified word count to approximately stop at
    Return: a list containing each paragraph containing approximate word_count
    '''

    # first split the text into individual words
    # iterate through until reaching the word_count OR reaching a null
    paragraphs = []

    count = 0
    current_paragraph = ''
    for word in text.split():
        current_paragraph += word + ' ' 
        count += 1

        # if word count is reached, then we want to keep going until reaching a dot
        if count >= word_count:

            # if there is a dot in the word, then end of sentence is reached.
            if "." in word:
                paragraphs.append(current_paragraph)
                count = 0
                current_paragraph = ''
    
    # add leftovers to the last paragraph
    paragraphs.append(current_paragraph)
        
    return paragraphs

def retrieveText(url):
    '''
    Retrieves text from a given website (in this case novelHall)
    website: a string URL to retrieve the text from
    '''
    driver.get(url)

    # find the text in the website
    elements = driver.find_elements(By.XPATH,'//*[@id="htmlContent"]')
    texts = [element.text for element in elements]

    # print extracted text and put into a text variable
    paragraph = ""
    for text in texts:
        paragraph += text

    driver.quit()   

    return paragraph

    
def addDoc(filename, header, paragraphs):
    # new document
    document = docx.Document()

    document.add_heading(header,0)

    for i in range(len(paragraphs)):
        document.add_paragraph(paragraphs[i])

    document.save(filename)
    

def main():

    # retrieve text from website
    website = str(input("Give the website address: "))
    
    novel_text = retrieveText(website)
    #novel_text = novel_text.replace("\n","")

    paragraphs = split_text_into_paragraphs(novel_text,200)

    addDoc("Chapter775-5.docx","Chapter 775",paragraphs)

    print("End of Program.")

    driver.quit()
    
    

if __name__ == "__main__":
    main()